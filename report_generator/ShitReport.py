'''
Author     : S1g0day
Version    : 0.8.4
Creat time : 2024/5/24 09:29
Modification time: 2024/8/5 16:47
Introduce  : 便携式报告编写工具
'''

import os
import sys
import yaml
import time
import tempfile
import warnings
import tldextract
import pandas as pd
from docx import Document
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PyQt6.QtGui import QPixmap, QIcon
from PyQt6.QtCore import Qt
from PyQt6.QtWidgets import QApplication, QListView, QWidget, QLabel, QLineEdit, QComboBox, QPushButton, QVBoxLayout, QHBoxLayout, QFormLayout, QMessageBox, QScrollArea
warnings.filterwarnings("ignore", category=DeprecationWarning)

# 自定义异常，用于中断嵌套循环
class InsertionError(Exception):
    pass

class ReportGenerator(QWidget):
    def __init__(self):
        super().__init__()
        
        '''读取文件'''
        # 从 YAML 文件中获取默认值
        self.push_config = yaml.safe_load(open("config/config.yaml", "r", encoding="utf-8").read())
        # 加载模板文件
        self.doc = Document(self.push_config["template_path"])
        # 读取Excel文件
        self.vulnerability_names, self.vulnerabilities = self.read_vulnerabilities_from_excel(self.push_config["vulnerabilities_file"])
        self.Icp_domains, self.Icp_infos = self.read_Icp_from_excel(self.push_config["icp_info_file"])

        '''设置窗口图标等其他初始化'''
        self.setWindowIcon(QIcon('config/th.jpg'))
        self.init_ui()  # 初始化UI界面

    def init_ui(self):
        '''设置 GUI 组件的初始化代码'''
        self.labels = ["隐患编号:", "隐患名称:", "隐患URL:", "隐患类型:", "隐患级别:",
                       "预警级别:", "归属地市:", "单位类型:", "所属行业:", "单位名称:",
                       "网站名称:", "网站域名:", "网站IP:", "备案号:", "发现时间:",
                       "漏洞描述:", "漏洞危害:", "修复建议:", "证据截图:", "工信域名备案截图:", 
                       "备注:"]
        

        self.text_edits = [QLineEdit(self) for _ in range(15)]

        # 创建文本框用于隐患编号
        self.vulnerability_id_text_edit = self.text_edits[0]
        # 设置文本框的初始文本为生成的隐患编号
        self.vulnerability_id_text_edit.setText(self.generate_vulnerability_id())
        # 创建漏洞类型下拉框
        self.vulName_box = QComboBox(self)
        self.vulName_box.addItems(self.vulnerability_names)

        # 设置下拉框样式
        self.vulName_box.setFixedSize(200, 20)
        self.vulName_box.setView(QListView())  ##todo 下拉框样式
        self.vulName_box.setStyleSheet(
            # "QComboBox {border: 1px solid #000000;background-color: rgb(255, 255, 255);font-size:12px;padding-left:14px;}"
            "QComboBox QAbstractItemView {font-size:14px;}"      # 下拉文字大小
            "QComboBox QAbstractItemView::item {height: 30 px;padding-left:20px;}"  # 下拉文字宽高
            "QScrollBar:vertical {border: 2px solid grey;width: 20px;}")    # 下拉侧边栏宽高

        # 创建漏洞等级下拉框
        self.hazardLevel_box = QComboBox(self)
        self.hazardLevel_box.addItems(['高危', '中危', '低危'])

        # 设置下拉框样式
        self.hazardLevel_box.setFixedSize(70, 20)
        self.hazardLevel_box.setView(QListView())  ##todo 下拉框样式
        self.hazardLevel_box.setStyleSheet(
            "QComboBox QAbstractItemView {font-size:12px;}"      # 下拉文字大小
            "QComboBox QAbstractItemView::item {height: 30 px;padding-left:20px;}"  # 下拉文字宽高
            "QScrollBar:vertical {border: 2px solid grey;width: 20px;}")    # 下拉侧边栏宽高

        # 创建文本框用于预警级别
        self.alert_level_text_edit = self.text_edits[1]
        # self.alert_level_text_edit.setReadOnly(True)  # 只读
        # 当hazardLevel_box值改变时调用update_alert_level方法
        self.hazardLevel_box.currentIndexChanged.connect(self.update_alert_level)
        # 初始化预警级别
        self.update_alert_level()

        # 创建单位类型下拉框
        self.unitType_box = QComboBox(self)
        self.unitType_box.addItems(self.push_config["unitType"])

        # 设置下拉框样式
        self.unitType_box.setFixedSize(100, 20)
        self.unitType_box.setView(QListView())  ##todo 下拉框样式
        self.unitType_box.setStyleSheet(
            "QComboBox QAbstractItemView {font-size:14px;}"      # 下拉文字大小
            "QComboBox QAbstractItemView::item {height: 30 px;padding-left:20px;}"  # 下拉文字宽高
            "QScrollBar:vertical {border: 2px solid grey;width: 20px;}")    # 下拉侧边栏宽高

        # 创建所属行业下拉框
        self.industry_box = QComboBox(self)
        self.industry_box.addItems(self.push_config["industry"])

        # 设置下拉框样式
        self.industry_box.setFixedSize(100, 20)
        self.industry_box.setView(QListView())  ##todo 下拉框样式
        self.industry_box.setStyleSheet(
            "QComboBox QAbstractItemView {font-size:12px;}"      # 下拉文字大小
            "QComboBox QAbstractItemView::item {height: 30 px;padding-left:20px;}"  # 下拉文字宽高
            "QScrollBar:vertical {border: 2px solid grey;width: 20px;}")    # 下拉侧边栏宽高

        # 创建文本框用于发现时间
        self.discovery_date_edit = self.text_edits[13]
        # 设置文本框的初始文本为当前日期
        current_date = datetime.now().strftime('%Y.%m.%d')
        self.discovery_date_edit.setText(current_date)

        # 创建用于显示工信域名备案截图的标签和按钮
        self.image_label_asset = QLabel(self)
        self.paste_button_asset = QPushButton('点击读取截图', self)
        self.paste_button_asset.clicked.connect(self.paste_image)
        self.delete_button_asset = QPushButton('删除图片', self)
        self.delete_button_asset.clicked.connect(self.delete_image)

        # 保存所有的漏洞复现描述部分
        self.vuln_sections = []
        # 添加按钮用于在界面上添加新的漏洞复现描述和漏洞证明图片的功能
        self.add_vuln_button = QPushButton('添加证明', self)
        self.add_vuln_button.clicked.connect(self.add_vulnerability_section)

        self.generate_button = QPushButton('生成报告', self)
        self.generate_button.clicked.connect(self.generate_report)

        self.reset_button = QPushButton('一键重置', self)
        self.reset_button.clicked.connect(self.reset_all)

        self.clear_all_button = QPushButton('一键清除', self)
        self.clear_all_button.clicked.connect(self.clear_all_sections)

        '''设置 GUI 组件表单布局'''
        self.form_layout = QFormLayout()
        # 添加用于隐患编号的文本框到布局
        self.form_layout.addRow(QLabel(self.labels[0]), self.vulnerability_id_text_edit)

        # 创建一个水平布局用于放置漏洞类型和漏洞等级
        h_layout = QHBoxLayout()
        # 添加漏洞类型下拉框到表单布局
        h_layout.addWidget(self.vulName_box)
        # 添加漏洞等级下拉框到表单布局
        h_layout.addWidget(QLabel(self.labels[4]))
        h_layout.addWidget(self.hazardLevel_box)
        # 添加自动更新预警级别到表单布局
        h_layout.addWidget(QLabel(self.labels[5]))
        h_layout.addWidget(self.alert_level_text_edit)
        self.form_layout.addRow(QLabel(self.labels[3]), h_layout)

        # 添加隐患URL到表单布局
        self.form_layout.addRow(QLabel(self.labels[2]), self.text_edits[5])
        self.text_edits[5].textChanged.connect(self.update_get_domain)
       
        # 创建一个水平布局用于放置域名信息
        Website_Name_layout = QHBoxLayout()
        # 添加网站名称到表单布局
        Website_Name_layout.addWidget(self.text_edits[6])
        # 添加网站IP到表单布局
        Website_Name_layout.addWidget(QLabel(self.labels[12]))
        Website_Name_layout.addWidget(self.text_edits[8])
        self.form_layout.addRow(QLabel(self.labels[10]), Website_Name_layout)

        # 创建一个水平布局用于放置域名信息
        domain_layout = QHBoxLayout()
        # 添加网站域名到表单布局
        domain_layout.addWidget(self.text_edits[7])
        # 添加工信备案号到表单布局
        domain_layout.addWidget(QLabel(self.labels[13]))
        domain_layout.addWidget(self.text_edits[9])
        self.form_layout.addRow(QLabel(self.labels[11]), domain_layout)
        # 添加单位名称到表单布局
        self.form_layout.addRow(QLabel(self.labels[9]), self.text_edits[3])

        # 在init_ui方法中，为网站域名添加信号，根据域名自动从文件中提取备案信息
        self.text_edits[7].textChanged.connect(self.update_icp_info)
        self.update_icp_info()

        # 在init_ui方法中，为单位名称、网站名称和隐患类型添加信号，根据其中变化自动调整其他参数数据
        self.text_edits[3].textChanged.connect(self.update_hazard_name)
        self.text_edits[6].textChanged.connect(self.update_hazard_name)
        self.vulName_box.currentIndexChanged.connect(self.update_hazard_name)
        self.update_hazard_name()

        # 创建一个水平布局用于放置公司信息
        unit_layout = QHBoxLayout()
        # 添加归属城市到表单布局
        self.text_edits[4].setText(self.push_config["city"])  # 设置默认值
        unit_layout.addWidget(self.text_edits[4])
        # 添加单位类型到表单布局
        unit_layout.addWidget(QLabel(self.labels[7]))
        unit_layout.addWidget(self.unitType_box)
        # 添加所属行业到表单布局
        unit_layout.addWidget(QLabel(self.labels[8]))
        unit_layout.addWidget(self.industry_box)
        self.form_layout.addRow(QLabel(self.labels[6]), unit_layout)

        # 添加隐患名称到表单布局
        self.form_layout.addRow(QLabel(self.labels[1]), self.text_edits[2])

        # 添加发现时间到表单布局
        self.form_layout.addRow(QLabel(self.labels[14]), self.discovery_date_edit)

        # 添加漏洞描述到表单布局
        self.form_layout.addRow(QLabel(self.labels[15]), self.text_edits[10])
        self.text_edits[10].setFixedHeight(60)  # 漏洞描述可能较长，增加文本框高度
        # 添加漏洞危害到表单布局
        self.form_layout.addRow(QLabel(self.labels[16]), self.text_edits[11])
        self.text_edits[11].setFixedHeight(60)  # 漏洞危害可能较长，增加文本框高度
        self.text_edits[11].textChanged.connect(self.update_hazard_name)
        # 添加整改建议到表单布局
        self.form_layout.addRow(QLabel(self.labels[17]), self.text_edits[12])
        self.text_edits[12].setFixedHeight(60)  # 整改建议可能较长，增加文本框高度

        # 添加备注到表单布局
        self.form_layout.addRow(QLabel(self.labels[-1]), self.text_edits[14])

        # 创建按钮布局用于放置生成报告
        button_layout = QHBoxLayout()
        button_layout.addWidget(self.generate_button)
        button_layout.addWidget(self.reset_button)  # 一键重置
        self.form_layout.addRow(button_layout)

        # 添加新的漏洞复现描述和图片按钮
        vuln_button_layout = QHBoxLayout()
        vuln_button_layout.addWidget(self.add_vuln_button)
        vuln_button_layout.addWidget(self.clear_all_button)
        self.form_layout.addRow(vuln_button_layout)

        # 添加工信域名备案截图到表单布局
        asset_layout = QHBoxLayout()
        asset_layout.addWidget(self.image_label_asset)
        asset_layout.addWidget(self.paste_button_asset)
        asset_layout.addWidget(self.delete_button_asset)
        self.form_layout.addRow(QLabel(self.labels[19]), asset_layout)

        # 调用 add_vulnerability_section 方法
        self.add_vulnerability_section()

        '''把表单布局添加到主布局中'''
        # 创建一个垂直布局，用于管理其他小部件和布局
        v_layout = QVBoxLayout()

        # 创建一个滚动区域，用于容纳可能超出屏幕显示范围的内容
        v_scroll = QScrollArea()

        # 将表单布局添加到垂直布局中
        v_layout.addLayout(self.form_layout)

        # 创建一个QWidget作为滚动区域的子部件
        widget = QWidget()

        # 将垂直布局设置为widget的布局
        widget.setLayout(v_layout)

        # 将widget设置为滚动区域的子部件
        v_scroll.setWidget(widget)

        # 设置滚动区域可以自动调整大小以适应其内容
        v_scroll.setWidgetResizable(True)

        # 指定滚动区域的宽度
        v_scroll.setFixedWidth(600)  # 或者使用 setMinimumWidth 根据需要
        # 创建主布局，用于管理整个窗口的内容
        main_layout = QVBoxLayout()

        # 将滚动区域添加到主布局中
        main_layout.addWidget(v_scroll)

        # 将主布局设置为QMainWindow的布局
        self.setLayout(main_layout)
        
        # 设置窗口的标题
        self.setWindowTitle(f'风险隐患报告生成器 - {self.push_config["version"]}')

        # 设置窗口的固定大小
        self.setFixedSize(620, 700)

        # 显示窗口
        self.show()

    def reset_all(self):
        self.vulnerability_id_text_edit.setText(self.generate_vulnerability_id())
        self.vulName_box.setCurrentIndex(0)
        self.hazardLevel_box.setCurrentIndex(0)
        # self.alert_level_text_edit.clear()    # 等级不需要清除
        self.unitType_box.setCurrentIndex(0)
        self.industry_box.setCurrentIndex(0)
        self.text_edits[2].clear()
        self.text_edits[3].clear()
        # self.text_edits[4].clear()    # 城市不需要清除
        self.text_edits[5].clear()
        self.text_edits[6].clear()
        self.text_edits[7].clear()
        self.text_edits[8].clear()
        self.text_edits[9].clear()
        self.text_edits[10].clear()
        self.text_edits[11].clear()
        self.text_edits[12].clear()
        self.text_edits[13].setText(datetime.now().strftime('%Y.%m.%d'))
        self.text_edits[14].clear()
        self.delete_image()
        self.clear_all_sections()

    def clear_all_sections(self):
        for section in self.vuln_sections:
            layout, edit, image_label = section
            self.form_layout.removeRow(layout)
            
        self.vuln_sections.clear()

    def update_get_domain(self):
        '''提取隐患url的根域名'''
        url = self.text_edits[5].text().strip()
        domain = tldextract.extract(url).registered_domain
        self.text_edits[7].setText(domain)  # 设置网站域名

    def update_icp_info(self):
        '''
        根据域名自动识别ICP备案信息
        '''
        domain = self.text_edits[7].text().strip()
        unit_name, service_licence = self.get_Icp_info(domain)

        self.text_edits[3].setText(unit_name)
        self.text_edits[9].setText(service_licence)

    # 添加一个槽函数用于更新隐患名称的值
    def update_hazard_name(self):
        unit_name = self.text_edits[3].text().strip()
        website_name = self.text_edits[6].text().strip()
        Vulnerability_Hazard = self.text_edits[11].text().strip()
        hazard_type = self.vulName_box.currentText()
        hazard_name = f"{unit_name}{website_name}存在{hazard_type}漏洞隐患"
        self.text_edits[2].setText(hazard_name)  # 设置隐患名称

        description, solution = self.get_vulnerability_info(hazard_type)
        # self.text_edits[10].setText(hazard_name if (not description or description == "无") else description)  
        
        # 设置漏洞描述
        if not description or description == "无":
            
            if len(Vulnerability_Hazard) == 0:    # 设置漏洞危害
                self.text_edits[10].setText(hazard_name)
            else:
                self.text_edits[10].setText(f"{hazard_name}{Vulnerability_Hazard}")
        else:
            if len(Vulnerability_Hazard) == 0:
                self.text_edits[10].setText(description)
            else:
                self.text_edits[10].setText(f"{description}{Vulnerability_Hazard}")

        self.text_edits[12].setText(solution)  # 设置整改建议

    def update_alert_level(self):
        # 更新预警级别
        hazard_level = self.hazardLevel_box.currentText()
        alert_level_map = {
            '高危': '3级',
            '中危': '4级',
            '低危': '4级'
        }
        alert_level = alert_level_map.get(hazard_level, '')
        self.alert_level_text_edit.setText(alert_level)

    def generate_vulnerability_id(self):
        # 根据系统时间生成隐患编号
        current_time = datetime.now().strftime('%Y-%m-%d-%H%M%S')
        return f"HN-XX-XX-{current_time}"

    def add_vulnerability_section(self):
        # 创建一个新的水平布局用于漏洞复现描述和相关操作按钮
        new_vuln_layout = QHBoxLayout()

        # 创建编辑框、标签和按钮
        new_vuln_edit = QLineEdit(self)
        new_vuln_image_label = QLabel(self)
        new_paste_button = QPushButton('点击读取截图', self)
        new_paste_button.clicked.connect(lambda: self.paste_new_vuln_image(new_vuln_image_label))
        new_delete_button = QPushButton('删除图片', self)
        new_delete_button.clicked.connect(lambda: self.delete_new_vuln_image(new_vuln_image_label))
        delete_section_button = QPushButton('删除该段', self)
        delete_section_button.clicked.connect(lambda: self.delete_vulnerability_section(new_vuln_layout, new_vuln_edit, new_vuln_image_label))

        # 将部件添加到新的水平布局中
        new_vuln_layout.addWidget(QLabel("漏洞复现描述:"))
        new_vuln_layout.addWidget(new_vuln_edit)
        new_vuln_layout.addWidget(new_vuln_image_label)
        new_vuln_layout.addWidget(new_paste_button)
        new_vuln_layout.addWidget(new_delete_button)
        new_vuln_layout.addWidget(delete_section_button)

        self.form_layout.addRow(new_vuln_layout)

        # 保存漏洞复现描述和图片路径
        self.vuln_sections.append((new_vuln_layout, new_vuln_edit, new_vuln_image_label))

    '''监控剪贴板'''
    def get_screenshot_from_clipboard(self):
        clipboard = QApplication.clipboard()
        mime_data = clipboard.mimeData()
        # 检查剪贴板数据是否为图片类型
        if mime_data.hasImage():
            # 获取 QImage 对象而不是原始数据
            image = clipboard.image()
            return image
        else:
            QMessageBox.warning(self, '错误', '剪贴板中没有图片！')
            return None

    '''处理漏洞复现截图'''
    def paste_new_vuln_image(self, image_label):
        screenshot = self.get_screenshot_from_clipboard()
        if screenshot:
            # pyqt6: 在 GUI 中显示缩放后的图片
            image_label.setPixmap(QPixmap.fromImage(screenshot).scaled(50, 50, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation))
            # 保存原始大小图片的引用
            image_label.original_pixmap = QPixmap.fromImage(screenshot)

    def delete_new_vuln_image(self, image_label):
        image_label.clear()

    def delete_vulnerability_section(self, layout, edit, label):
        for i in reversed(range(layout.count())):
            widget = layout.itemAt(i).widget()
            if widget is not None:
                widget.setParent(None)
        self.form_layout.removeRow(layout)
        self.vuln_sections.remove((layout, edit, label))

    '''处理备案截图'''
    def paste_image(self, image_type):
        """粘贴图像到 QLabel 并保存图像路径"""
        screenshot = self.get_screenshot_from_clipboard()
        if screenshot:
            self.asset_image = screenshot
            # pyqt6: 在 GUI 中显示缩放后的图片
            self.image_label_asset.setPixmap(QPixmap.fromImage(screenshot).scaled(50, 50, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation))

    # 删除图片的函数
    def delete_image(self):
        """删除 QLabel 中的图像"""
        self.image_label_asset.clear()

    def replace_report_text(self, replacements):
        # 替换段落中的占位符
        for paragraph in self.doc.paragraphs:
            runs = paragraph.runs
            for i, run in enumerate(runs):
                if run.text == '#':
                    counter = i  # 记录起始位置
                    tmp = '#'    # tmp开始存储
                    while tmp not in list(replacements.keys()):
                        counter += 1
                        tmp += runs[counter].text
                        runs[counter].clear()
                    runs[i].text = runs[i].text.replace(runs[i].text, replacements[tmp])

        # 替换表格中的占位符
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        runs = paragraph.runs
                        full_text = ''.join(run.text for run in runs)

                        for key, value in replacements.items():
                            if key in full_text:
                                remaining_text = full_text
                                for run in runs:
                                    if key in remaining_text:
                                        start_index = remaining_text.index(key)
                                        end_index = start_index + len(key)

                                        pre_key_text = remaining_text[:start_index]
                                        post_key_text = remaining_text[end_index:]

                                        run.text = pre_key_text + value
                                        remaining_text = post_key_text
                                    else:
                                        run.text = remaining_text
                                        remaining_text = ''

    # 处理单图文
    def replace_text_with_image(self, key, img_path):
        previous_content = None  # 用于跟踪上一个打印的内容
        for table_index, table in enumerate(self.doc.tables):
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    content = cell.text.strip()
                    if content != previous_content:
                        previous_content = content  # 更新上一个打印的内容
                        if key in content:
                            # 仅删除关键字
                            cell.text = cell.text.replace(key, "")
                            # 在现有段落中插入图片，并设置图片居中
                            paragraph = cell.paragraphs[0]
                            run = paragraph.add_run()
                            # 获取单元格宽度
                            cell_width = cell.width
                            run.add_picture(img_path, width=cell_width)  # 设置合适的宽度

                            # 将刚刚插入的图片居中
                            last_run = paragraph.runs[-1]
                            last_run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 处理多图片
    def vuln_sections_text_with_image(self, vuln_text, img_path):
        previous_content = None  # 用于跟踪上一个打印的内容
        try:
            for table_index, table in enumerate(self.doc.tables):
                for row_index, row in enumerate(table.rows):
                    for cell_index, cell in enumerate(row.cells):
                        content = cell.text.strip()
                        if content != previous_content:
                            previous_content = content  # 更新上一个打印的内容
                            # 获取表格中的第12行第0列的单元格的段落
                            paragraph = self.doc.tables[0].cell(13, 0).paragraphs[0]
                            # 在段落中添加一个文本运行，用于显示漏洞文本
                            if vuln_text == '':
                                paragraph.add_run(vuln_text)
                            else:
                                paragraph.add_run(vuln_text + '\n')
                            # 在段落中添加一个空的文本运行
                            run = paragraph.add_run()
                            # 获取单元格的宽度
                            cell_width = cell.width
                            # 在段落中添加一个图片运行，并设置宽度为单元格的宽度
                            run.add_picture(img_path, width=cell_width)
                            # 获取段落中最后一个运行
                            last_run = paragraph.runs[-1]
                            # 将最后一个运行的对齐方式设置为居中
                            last_run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            raise InsertionError  # 成功插入后抛出异常，中断所有循环

        except InsertionError:
            pass  # 捕获异常以防止程序崩溃，并正常继续运行

    def process_vuln_sections(self):
        for layout, edit, label in self.vuln_sections:
            if hasattr(label, 'original_pixmap'):  # 检查是否有原始图片的引用
                # 使用原始图片的路径
                vuln_path = self.save_image_temporarily(label.original_pixmap.toImage())
                self.vuln_sections_text_with_image(edit.text().strip(), vuln_path)

    def save_image_temporarily(self, image):
        temp_file = tempfile.mkstemp(suffix='.png')[1]
        image.save(temp_file)
        return temp_file

    def save_document(self, report_file_path_notime):
        if not os.path.exists(report_file_path_notime):
            # 文件不存在，直接保存文档
            self.doc.save(report_file_path_notime)
            return report_file_path_notime
        else:
            # 文件已存在
            count = 1
            while os.path.exists(f'{report_file_path_notime[:-5]}-{count}.docx'):
                # 根据规则生成新的文件名，继续检查是否存在
                count += 1
            new_file_path = f'{report_file_path_notime[:-5]}-{count}.docx'
            self.doc.save(new_file_path)
            return new_file_path
        
    def log_save(self, replacements):

        # 获取脚本所在目录
        script_dir = os.path.dirname(__file__)

        # 创建日志目录
        log_dir = os.path.join(script_dir, 'output/')
        if not os.path.exists(log_dir):
            os.makedirs(log_dir)

        # 获取当前时间戳
        timestamp = int(time.time())

        # 创建客户公司目录
        customerCompanyName_dir = os.path.join(script_dir, f'output/{replacements["#customerCompanyName#"]}')
        if not os.path.exists(customerCompanyName_dir):
            os.makedirs(customerCompanyName_dir)

        # 构建报告文件路径
        # report_file_path = f'{customerCompanyName_dir}/{timestamp}_{replacements["#customerCompanyName#"]}{replacements["#websitename#"]}存在{replacements["#vulName#"]}漏洞隐患【{replacements["#hazardLevel#"]}】.docx'
        report_file_path_notime = f'{customerCompanyName_dir}/{replacements["#customerCompanyName#"]}{replacements["#websitename#"]}存在{replacements["#vulName#"]}漏洞隐患【{replacements["#hazardLevel#"]}】.docx'      
        # 保存文档
        report_file_path = self.save_document(report_file_path_notime)
        # 显示一个消息框通知用户报告已生成
        QMessageBox.information(self, '报告生成', f'报告已生成: {report_file_path}')

        output_file = f'{replacements["#customerCompanyName#"]}\t{replacements["#target#"]}\t{replacements["#vulName#"]}\t{self.push_config["supplierName"]}\t{replacements["#reportTime#"]}'
        output_file_path = f'output/{replacements["#reportTime#"]}_output.txt'
        with open(output_file_path, 'a+') as f: f.write('\n'+output_file)

    def read_vulnerabilities_from_excel(self, file_path):
        """
        从Excel文件中读取漏洞信息，并将其存储到字典中
        """
        data = pd.read_excel(file_path)
        
        vulnerability_names = []  # 存储漏洞名称的列表
        vulnerabilities = {}  # 存储漏洞描述和加固建议的字典
        
        for index, row in data.iterrows():
            vulnerability_name = row['漏洞名称']
            vulnerability_description = row['漏洞描述']
            vulnerability_solution = row['加固建议']
            
            vulnerability_names.append(vulnerability_name)  # 将漏洞名称添加到列表中
            
            vulnerabilities[vulnerability_name.lower()] = {
                '漏洞描述': vulnerability_description,
                '加固建议': vulnerability_solution
            }
        
        return vulnerability_names, vulnerabilities
    
    def get_vulnerability_info(self, name):
        """
        根据漏洞名称从字典中获取漏洞描述和加固建议
        """
        name = name.lower()
        if name in self.vulnerabilities:
            description = self.vulnerabilities[name]['漏洞描述']
            solution = self.vulnerabilities[name]['加固建议']
            return description, solution
        else:
            return None, None

    def read_Icp_from_excel(self, file_path):
        """
        从Excel文件中读取ICP信息，并将其存储到字典中
        """
        data = pd.read_excel(file_path)
        Icp_domains = []  # 存储根域名的列表
        Icp_infos = {}  # 存储单位名称和备案号的字典
        
        for index, row in data.iterrows():
            Icp_domain = row['domain']
            Icp_serviceLicence = row['serviceLicence']
            Icp_unitName = row['unitName']
            
            Icp_domains.append(Icp_domain)  # 将根域名添加到列表中
            
            Icp_infos[Icp_domain.lower()] = {
                'serviceLicence': Icp_serviceLicence,
                'unitName': Icp_unitName
            }
        
        return Icp_domains, Icp_infos
    
    def get_Icp_info(self, domain):
        """
        根据根域名从字典中获取单位名称和备案号
        """
        name = domain.lower()
        if name in self.Icp_infos:
            unitName = self.Icp_infos[name]['unitName']
            serviceLicence = self.Icp_infos[name]['serviceLicence']
            return unitName, serviceLicence
        else:
            return None, None

    '''主函数'''
    def generate_report(self):

        # 创建一个字典，包含所有需要替换的字段
        replacements = {
            '#reportId#': self.text_edits[0].text().strip(),
            '#reportName#': self.text_edits[2].text().strip(),
            '#target#': self.text_edits[5].text().strip(),
            '#vulName#': self.vulName_box.currentText(),
            '#hazardLevel#': self.hazardLevel_box.currentText(),
            '#warningLevel#': self.alert_level_text_edit.text().strip(),
            '#city#': self.text_edits[4].text().strip(),
            '#unitType#': self.unitType_box.currentText(),
            '#industry#': self.industry_box.currentText(),
            '#customerCompanyName#': self.text_edits[3].text().strip(),
            '#websitename#': self.text_edits[6].text().strip(),
            '#domain#': self.text_edits[7].text().strip(),
            '#ipaddress#': self.text_edits[8].text().strip(),
            '#caseNumber#': self.text_edits[9].text().strip(),
            '#reportTime#': self.discovery_date_edit.text().strip(),
            '#problemDescription#': self.text_edits[10].text().strip(),
            '#vul_modify_repair#': self.text_edits[12].text().strip(),
            '#remark#': self.text_edits[14].text().strip(),
        }
        self.replace_report_text(replacements)

        # 添加工信域名备案截图
        if hasattr(self, 'asset_image') and self.asset_image:
            asset_path = self.save_image_temporarily(self.asset_image)
            self.replace_text_with_image("#screenshotoffiling#", asset_path)

        # 处理单个或多个漏洞复现描述和图片
        self.process_vuln_sections()

        # 保存日志及文件
        self.log_save(replacements)

if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = ReportGenerator()
    ex.show()
    sys.exit(app.exec())
