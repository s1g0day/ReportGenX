import tempfile
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class DocumentImageProcessor:
    def __init__(self, doc, vuln_sections):
        self.doc = doc
        self.vuln_sections = vuln_sections

    def text_with_image(self, content, img_path):
        previous_content = None  # 用于跟踪上一个打印的内容
        for _, table in enumerate(self.doc.tables):
            for _, row in enumerate(table.rows):
                for _, cell in enumerate(row.cells):
                    cell_content = cell.text.strip()
                    if cell_content != previous_content:
                        previous_content = cell_content  # 更新上一个打印的内容
                        if content != "#screenshotoffiling#":
                            # 获取表格中的第12行第0列的单元格的段落
                            paragraph = self.doc.tables[0].cell(13, 0).paragraphs[0]
                            # 在段落中添加一个文本运行，用于显示漏洞文本
                            if content == '':
                                paragraph.add_run(content)
                            else:
                                paragraph.add_run(content + '\n')
                            # 在段落中添加一个空的文本运行
                            run = paragraph.add_run()
                            # 获取单元格的宽度
                            cell_width = cell.width
                            # 在段落中添加一个图片运行，并设置宽度为单元格的宽度
                            run.add_picture(img_path, width=cell_width)
                            # 获取段落中最后一个运行
                            last_run = paragraph.runs[-1]
                            # 将最后一个运行的对齐方式设置为居中
                            last_run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            return #成功插入漏洞文本和图片后，代码会立即返回，从而中断所有循环
                        else:
                            if content in cell_content:
                                # 仅删除关键字
                                cell.text = cell.text.replace(content, "")
                                # 在现有段落中插入图片，并设置图片居中
                                paragraph = cell.paragraphs[0]
                                run = paragraph.add_run()
                                # 获取单元格宽度
                                cell_width = cell.width
                                run.add_picture(img_path, width=cell_width)  # 设置合适的宽度

                                # 将刚刚插入的图片居中
                                last_run = paragraph.runs[-1]
                                last_run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def save_image_temporarily(self, image):
        temp_file = tempfile.mkstemp(suffix='.png')[1]
        image.save(temp_file)
        return temp_file
    
    def process_vuln_sections(self):
        for _, edit, label in self.vuln_sections:
            if hasattr(label, 'original_pixmap'):  # 检查是否有原始图片的引用
                # 使用原始图片的路径
                vuln_path = self.save_image_temporarily(label.original_pixmap.toImage())
                self.text_with_image(edit.text().strip(), vuln_path)
            
